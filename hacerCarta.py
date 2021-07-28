from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Titulo del documento', 0) # Rodrigo : Es el titulo que presentara el documento

p = document.add_paragraph('Un parrafo con ')
p.add_run('bold').bold = True   # Rodrigo : Se le agrega Negrita
p.add_run(' y con ')
p.add_run('italic.').italic = True # Rodrigo : Se le agrega Cursiva

document.add_heading('Subtitulo 1', level=1)   # Rodrigo : Como agregar un subtitulo 
document.add_paragraph('Sub tipo 1') # Rodrigo : Como agregar un parrafo 

document.add_paragraph(
    'primer elemento de la lista desordenada ', style='List Bullet' # Rodrigo : Se le agrega el punto negro al inicio del parrafo
)
document.add_paragraph(
    'primer elemento de la lista desordenada ', style='List Number' # Rodrigo : Se le agrega un numero al inicio del parrafo
)

records = (
    (3, '101', 'Cosas'),
    (7, '422', 'Huevos'),
    (4, '631', 'cosas, cosas, huevos, and cosas')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'CTD'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('Prueba.docx')    # Rodrigo : Nombre final al documento al momento de guardar